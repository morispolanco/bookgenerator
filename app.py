import streamlit as st
import requests
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
import re

# Función para limpiar Markdown
def clean_markdown(text):
    """Elimina marcas de Markdown del texto."""
    text = re.sub(r'[#*_`]', '', text)  # Eliminar caracteres especiales de Markdown
    return text.strip()

# Función para procesar listas y diálogos, reemplazando guiones por rayas
def process_dialogues_and_lists(text):
    """
    Procesa el texto para:
    1. Reemplazar guiones ('-') al inicio de las listas o diálogos por rayas ('—').
    2. Asegurar que después de las listas haya un salto de párrafo.
    """
    lines = text.split('\n')  # Dividir el texto en líneas
    processed_lines = []
    in_list = False  # Indicador para saber si estamos dentro de una lista o diálogo

    for line in lines:
        stripped_line = line.strip()
        if stripped_line.startswith('-'):  # Detectar líneas que comienzan con un guion
            # Reemplazar el guion por una raya
            processed_line = stripped_line.replace('-', '—', 1)
            processed_lines.append(processed_line)
            in_list = True
        else:
            if in_list:
                # Si salimos de una lista o diálogo, añadir un salto de párrafo
                processed_lines.append("")  # Salto de párrafo
                in_list = False
            processed_lines.append(stripped_line)

    # Unir las líneas procesadas con saltos de párrafo
    return '\n\n'.join(processed_lines)

# Función para aplicar reglas de capitalización según el idioma
def format_title(title, language):
    """
    Formatea el título según las reglas gramaticales del idioma.
    - Español: Solo mayúscula inicial en la primera palabra y nombres propios.
    - Otros idiomas: Mayúscula inicial en cada palabra.
    """
    if language.lower() == "spanish":
        words = title.split()
        formatted_words = [words[0].capitalize()] + [word.lower() for word in words[1:]]
        return " ".join(formatted_words)
    else:
        return title.title()

# Función para generar un capítulo usando Google Gemini
def generate_chapter(api_key, topic, audience, chapter_number, language, table_of_contents="", specific_instructions="", is_intro=False, is_conclusion=False):
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={api_key}"
    
    # Construir el mensaje con la tabla de contenido e instrucciones específicas
    if is_intro:
        message_content = f"Escribe la introducción sobre {topic} dirigida a {audience}."
    elif is_conclusion:
        message_content = f"Escribe las conclusiones sobre {topic} dirigidas a {audience}."
    else:
        message_content = f"Escribe el capítulo {chapter_number} sobre {topic} dirigido a {audience}."
    
    if table_of_contents:
        message_content += f" Sigue esta estructura: {table_of_contents}"
    
    if specific_instructions:
        message_content += f" {specific_instructions}"
    
    data = {
        "contents": [{"role": "user", "parts": [{"text": message_content}]}],
        "generationConfig": {
            "temperature": 1,
            "topK": 40,
            "topP": 0.95,
            "responseMimeType": "text/plain"
        }
    }
    
    try:
        response = requests.post(url, json=data)
        response.raise_for_status()
        content = response.json().get("candidates", [{}])[0].get("content", {}).get("parts", [{}])[0].get("text", "Error generating the chapter.")
    except Exception as e:
        st.error(f"Error generating chapter {chapter_number}: {str(e)}")
        content = "Error generating the chapter."
    
    # Procesar diálogos y listas
    processed_content = process_dialogues_and_lists(content)
    
    return clean_markdown(processed_content)

# Función para agregar numeración de páginas al documento Word
def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)

# Función para crear un documento Word con formato específico
def create_word_document(chapters, title, author_name, author_bio, language):
    doc = Document()

    # Configurar el tamaño de página (5.5 x 8.5 pulgadas)
    section = doc.sections[0]
    section.page_width = Inches(5.5)
    section.page_height = Inches(8.5)

    # Configurar márgenes de 0.8 pulgadas en todo
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    # Añadir título formateado según el idioma
    formatted_title = format_title(title, language)
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(formatted_title)
    title_run.bold = True
    title_run.font.size = Pt(14)
    title_run.font.name = "Times New Roman"

    # Añadir nombre del autor si está proporcionado
    if author_name:
        author_paragraph = doc.add_paragraph()
        author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_paragraph.add_run(author_name)
        author_run.font.size = Pt(12)
        author_run.font.name = "Times New Roman"
        doc.add_page_break()

    # Añadir perfil del autor si está proporcionado
    if author_bio:
        bio_paragraph = doc.add_paragraph("Author Information")
        bio_paragraph.style = "Heading 2"
        bio_paragraph.runs[0].font.size = Pt(11)
        bio_paragraph.runs[0].font.name = "Times New Roman"
        doc.add_paragraph(author_bio).style = "Normal"
        doc.add_page_break()

    # Añadir capítulos
    for i, chapter in enumerate(chapters, 1):
        chapter_title_text = f"Chapter {i}" if language.lower() != "spanish" else f"Capítulo {i}"
        formatted_chapter_title = format_title(chapter_title_text, language)
        chapter_title = doc.add_paragraph(formatted_chapter_title)
        chapter_title.style = "Heading 1"
        chapter_title.runs[0].font.size = Pt(12)
        chapter_title.runs[0].font.name = "Times New Roman"

        paragraphs = chapter.split('\n\n')
        for para_text in paragraphs:
            para_text = para_text.replace('\n', ' ').strip()
            paragraph = doc.add_paragraph(para_text)
            paragraph.style = "Normal"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(11)
                run.font.name = "Times New Roman"

        doc.add_page_break()

    add_page_numbers(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Configuración de Streamlit
st.set_page_config(page_title="Automatic Book Generator", page_icon="📚")

# Título con ícono
st.title("📚 Automatic Book Generator")

# Barra lateral con instrucciones y anuncio
st.sidebar.header("📖 How does this app work?")
st.sidebar.markdown("""
This application automatically generates books in `.docx` format based on a topic and target audience.  
The books can be **fiction** or **non-fiction**, depending on your input.  

**Steps to use it:**
1. Enter the book's topic.
2. Specify the target audience.
3. Provide an optional table of contents.
4. Write optional specific instructions.
5. Select the number of chapters desired (maximum 50).
6. Choose the book's language.
7. Decide whether to include an introduction, conclusions, author name, and author profile.
8. Click "Generate Book".
9. Download the generated file.
""")
st.sidebar.markdown("""
---
**📝 Text correction in 24 hours**  
👉 [Hablemos Bien](https://hablemosbien.org)
""")

# Validación de claves secretas
if "GOOGLE_API_KEY" not in st.secrets:
    st.error("Please configure the API key in Streamlit secrets.")
    st.stop()
api_key = st.secrets["GOOGLE_API_KEY"]

# Entradas del usuario
topic = st.text_input("📒 Book Topic:")
audience = st.text_input("🎯 Target Audience:")
table_of_contents = st.text_area("📚 Optional Table of Contents:", placeholder="Provide a table of contents for longer chapters.")
specific_instructions = st.text_area("📝 Optional Specific Instructions:", placeholder="Provide specific instructions for the book.")
num_chapters = st.slider("🔢 Number of Chapters", min_value=1, max_value=50, value=25)
include_intro = st.checkbox("Include Introduction", value=True)
include_conclusion = st.checkbox("Include Conclusions", value=True)
author_name = st.text_input("🖋️ Author Name (optional):")
author_bio = st.text_area("👤 Author Profile (optional):", placeholder="Brief professional description or biography.")
languages = ["English", "Spanish", "French", "German", "Chinese", "Japanese", "Russian", "Portuguese", "Italian", "Arabic", "Medieval Latin", "Koine Greek"]
selected_language = st.selectbox("🌐 Choose the book's language:", languages)

# Estado de Streamlit para almacenar los capítulos generados
if 'chapters' not in st.session_state:
    st.session_state.chapters = []

# Botón para generar el libro
if st.button("🚀 Generate Book"):
    if not topic or not audience:
        st.error("Please enter a valid topic and target audience.")
        st.stop()

    chapters = []

    # Generar introducción si está seleccionada
    if include_intro:
        st.write("⏳ Generating introduction...")
        intro_content = generate_chapter(api_key, topic, audience, 0, selected_language.lower(), table_of_contents, specific_instructions, is_intro=True)
        chapters.append(intro_content)
        word_count = len(intro_content.split())
        with st.expander(f"🌟 Introduction ({word_count} words)"):
            st.write(intro_content)

    # Generar capítulos principales
    progress_bar = st.progress(0)
    for i in range(1, num_chapters + 1):
        st.write(f"⏳ Generating chapter {i}...")
        chapter_content = generate_chapter(api_key, topic, audience, i, selected_language.lower(), table_of_contents, specific_instructions)
        word_count = len(chapter_content.split())  # Contar palabras
        chapters.append(chapter_content)
        with st.expander(f"📖 Chapter {i} ({word_count} words)"):
            st.write(chapter_content)
        progress_bar.progress(i / num_chapters)

    # Generar conclusiones si están seleccionadas
    if include_conclusion:
        st.write("⏳ Generating conclusions...")
        conclusion_content = generate_chapter(api_key, topic, audience, 0, selected_language.lower(), table_of_contents, specific_instructions, is_conclusion=True)
        word_count = len(conclusion_content.split())
        chapters.append(conclusion_content)
        with st.expander(f"🔚 Conclusions ({word_count} words)"):
            st.write(conclusion_content)

    st.session_state.chapters = chapters

# Mostrar opciones de descarga si hay capítulos generados
if st.session_state.chapters:
    st.subheader("⬇️ Download Options")
    word_file = create_word_document(st.session_state.chapters, topic, author_name, author_bio, selected_language.lower())

    st.download_button(
        label="📥 Download in Word",
        data=word_file,
        file_name=f"{topic}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
